"""
docx_extraction.py
------------------
Extract structured data from .docx files using:
    - python-docx : paragraphs, headings, tables, styles
    - mammoth     : clean HTML conversion (good for RAG chunking)

Input : ../data/docx/*.docx
Output:
    ../output/json/<name>_docx.json
    ../output/dataframes/<name>_docx_table<i>.csv   (one per table found)
    ../output/key_value/<name>_docx.txt

Install:
    pip install python-docx mammoth pandas
"""

import json
from pathlib import Path

import docx  # python-docx
import mammoth
import pandas as pd

BASE_DIR = Path(__file__).resolve().parent.parent
INPUT_DIR = BASE_DIR / "data" / "docx"
OUTPUT_JSON_DIR = BASE_DIR / "output" / "json"
OUTPUT_DF_DIR = BASE_DIR / "output" / "dataframes"
OUTPUT_KV_DIR = BASE_DIR / "output" / "key_value"

for d in (OUTPUT_JSON_DIR, OUTPUT_DF_DIR, OUTPUT_KV_DIR):
    d.mkdir(parents=True, exist_ok=True)


def extract_paragraphs(document: docx.Document):
    """Return list of paragraph dicts with text + style (heading level, etc.)."""
    paragraphs = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({
            "text": text,
            "style": para.style.name if para.style else None,
            "is_heading": para.style.name.lower().startswith("heading") if para.style else False,
        })
    return paragraphs


def extract_tables(document: docx.Document):
    """Return list of tables, each as list-of-rows (list of cell strings)."""
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return tables


def extract_key_value_pairs(paragraphs):
    """
    Heuristic key-value extraction: lines like 'Key: Value' or 'Key - Value'.
    Useful for forms / letterheads / metadata blocks.
    """
    kv_pairs = {}
    for p in paragraphs:
        text = p["text"]
        for sep in (":", " - ", "\t"):
            if sep in text:
                key, _, value = text.partition(sep)
                key, value = key.strip(), value.strip()
                if key and value and len(key) < 60:
                    kv_pairs[key] = value
                break
    return kv_pairs


def convert_with_mammoth(filepath: Path):
    """Use mammoth to get clean HTML (handles styles/images better for RAG chunking)."""
    with open(filepath, "rb") as f:
        result = mammoth.convert_to_html(f)
    html = result.value
    messages = [str(m) for m in result.messages]
    return html, messages


def process_docx_file(filepath: Path):
    document = docx.Document(str(filepath))

    paragraphs = extract_paragraphs(document)
    tables = extract_tables(document)
    key_values = extract_key_value_pairs(paragraphs)
    html, messages = convert_with_mammoth(filepath)

    structured = {
        "source_file": filepath.name,
        "num_paragraphs": len(paragraphs),
        "num_tables": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
        "key_value_pairs": key_values,
        "mammoth_html": html,
        "mammoth_messages": messages,
    }

    # ---- Save JSON ----
    json_path = OUTPUT_JSON_DIR / f"{filepath.stem}_docx.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(structured, f, indent=2, ensure_ascii=False)

    # ---- Save DataFrame(s) for each table ----
    for i, table in enumerate(tables):
        if not table:
            continue
        df = pd.DataFrame(table[1:], columns=table[0]) if len(table) > 1 else pd.DataFrame(table)
        df_path = OUTPUT_DF_DIR / f"{filepath.stem}_docx_table{i}.csv"
        df.to_csv(df_path, index=False)

    # ---- Save key-value text ----
    kv_path = OUTPUT_KV_DIR / f"{filepath.stem}_docx.txt"
    with open(kv_path, "w", encoding="utf-8") as f:
        for k, v in key_values.items():
            f.write(f"{k}: {v}\n")

    print(f"[OK] Processed {filepath.name} -> {json_path.name}")
    return structured


def main():
    docx_files = list(INPUT_DIR.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {INPUT_DIR}")
        return

    for filepath in docx_files:
        try:
            process_docx_file(filepath)
        except Exception as e:
            print(f"[ERROR] Failed to process {filepath.name}: {e}")


if __name__ == "__main__":
    main()